"""
generate_media_plan.py — Собирает артефакты пайплайна создания кампании в Директе
в единый медиаплан (docx).

Использование:
    python -m scripts.generate_media_plan --workspace <path/to/direct-campaigns/<slug>>
                                          [--output <path/to/output.docx>]

Ожидает в workspace md-файлы пайплайна (см. STEPS ниже) и опционально _state.json
с метаданными.

Если python-docx недоступен, фоллбечится на markdown-вариант — собирается один
media_plan.md.
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from datetime import date


# Артефакты нового пайплайна создания кампании в Директе.
# Опциональные файлы (например, 08_5_rsya_creatives_TODO.md) добавляются только если найдены.
STEPS = [
    ("00_brief", "0. Бриф"),
    ("01_masks_matrix", "1. Маски: матрица плюс-слов"),
    ("02_frequency", "2. Прогноз частотности и CPC"),
    ("04_wordstat_filter_summary", "4. Кластеризация семантики"),
    ("05_campaign_structure", "5. Структура кампаний"),
    ("06_negative_summary", "6. Минус-слова и кросс-минусация"),
    ("07_metrika_goals_plan", "7. Метрика и цели"),
    ("08_creatives", "8. Объявления"),
    ("08_5_rsya_creatives_TODO", "8.5. РСЯ-визуалы (TODO)"),
    ("09_bidding_strategy", "9. Стратегия торгов"),
    ("10_launch_log", "10. Лог заливки через MCP"),
]


def read_step(workspace: Path, step_id: str) -> str | None:
    """Читает md-файл шага. Возвращает None, если файла нет — для опциональных шагов."""
    path = workspace / f"{step_id}.md"
    if not path.exists():
        return None
    return path.read_text(encoding="utf-8")


def assemble_markdown(workspace: Path) -> str:
    state_path = workspace / "_state.json"
    state = {}
    if state_path.exists():
        try:
            state = json.loads(state_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass

    product = state.get("product_name", workspace.name)
    today = date.today().isoformat()

    out_parts = [
        f"# Медиаплан запуска в Яндекс.Директе: {product}",
        f"_Сгенерировано {today}_",
        "",
        "## Оглавление",
    ]
    available_steps = []
    for step_id, title in STEPS:
        content = read_step(workspace, step_id)
        if content is None and step_id == "08_5_rsya_creatives_TODO":
            # Опциональный шаг — пропускаем без отметки
            continue
        if content is None:
            # Обязательный шаг отсутствует — отмечаем
            content = f"_(шаг не заполнен: файл `{step_id}.md` не найден)_"
        anchor = title.split(". ", 1)[1].lower().replace(" ", "-").replace(":", "")
        out_parts.append(f"- [{title}](#{anchor})")
        available_steps.append((title, content))
    out_parts.append("")

    for title, content in available_steps:
        out_parts.append(f"## {title}")
        out_parts.append("")
        out_parts.append(content)
        out_parts.append("")

    return "\n".join(out_parts)


def write_docx(markdown_text: str, output_path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return False

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            cleaned = [
                row for row in table_rows
                if not all(re.fullmatch(r"\s*:?-+:?\s*", cell) for cell in row)
            ]
            if cleaned:
                table = doc.add_table(rows=len(cleaned), cols=len(cleaned[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(cleaned):
                    for c, cell in enumerate(row):
                        if c < len(cleaned[0]):
                            table.cell(r, c).text = cell.strip()
        table_rows = []
        in_table = False

    for line in lines:
        if line.startswith("| ") and "|" in line[2:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    if in_table:
        flush_table()

    doc.save(output_path)
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description="Собрать медиаплан Директа из артефактов пайплайна")
    parser.add_argument("--workspace", required=True, help="Папка direct-campaigns/<slug>")
    parser.add_argument("--output", default=None, help="Куда сохранить (по умолч. <workspace>/media_plan.docx)")
    args = parser.parse_args()

    workspace = Path(args.workspace).resolve()
    if not workspace.exists() or not workspace.is_dir():
        raise SystemExit(f"Workspace {workspace} не найден")

    output = Path(args.output) if args.output else workspace / "media_plan.docx"

    md = assemble_markdown(workspace)

    md_path = workspace / "media_plan.md"
    md_path.write_text(md, encoding="utf-8")
    print(f"Markdown собран: {md_path}")

    if write_docx(md, output):
        print(f"DOCX готов: {output}")
    else:
        print("python-docx не установлен — пропускаю генерацию docx. Установи: pip install python-docx --break-system-packages")


if __name__ == "__main__":
    main()
